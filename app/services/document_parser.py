"""
文档解析服务
支持PDF和DOCX格式，使用AI提取结构化内容
"""
import os
from pathlib import Path
from typing import List, Dict, Optional
import json
from datetime import datetime

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

from app.config import DOCUMENTS_DIR
from app.models.database import db


class DocumentParser:
    """文档解析器"""
    
    def __init__(self):
        self.supported_formats = {".pdf", ".docx", ".doc"}
    
    def parse_document(self, file_path: str, filename: str) -> Dict:
        """
        解析文档并提取内容
        
        Args:
            file_path: 文件路径
            filename: 文件名
            
        Returns:
            包含解析结果的字典
        """
        file_ext = Path(filename).suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {file_ext}")
        
        try:
            if file_ext == ".pdf":
                content = self._parse_pdf(file_path)
            elif file_ext in {".docx", ".doc"}:
                content = self._parse_docx(file_path)
            else:
                raise ValueError(f"未实现的解析器: {file_ext}")
            
            # 使用AI提取结构化内容（简化版，实际应使用AI模型）
            structured_content = self._extract_structured_content(content, filename)
            
            return {
                "raw_content": content,
                "structured_content": structured_content,
                "metadata": {
                    "filename": filename,
                    "format": file_ext,
                    "parsed_at": datetime.now().isoformat(),
                    "word_count": len(content.split()),
                    "char_count": len(content)
                }
            }
        except Exception as e:
            raise Exception(f"文档解析失败: {str(e)}")
    
    def _parse_pdf(self, file_path: str) -> str:
        """解析PDF文件"""
        if PdfReader is None:
            raise ImportError("pypdf未安装，请运行: pip install pypdf")
        
        try:
            reader = PdfReader(file_path)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"--- 第 {page_num + 1} 页 ---\n{text}")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"PDF解析错误: {str(e)}")
    
    def _parse_docx(self, file_path: str) -> str:
        """解析DOCX文件"""
        if Document is None:
            raise ImportError("python-docx未安装，请运行: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # 提取表格内容
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_data))
                if table_text:
                    text_parts.append("\n表格:\n" + "\n".join(table_text))
            
            return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"DOCX解析错误: {str(e)}")
    
    def _extract_structured_content(self, content: str, filename: str) -> Dict:
        """
        使用AI提取结构化内容
        这里是一个简化版本，实际应该使用AI模型（如GPT、Claude等）
        
        AI使用场景：
        - 提取关键信息（标题、章节、表格数据）
        - 识别文档类型和主题
        - 提取实体和关键词
        - 结构化表格数据（如HR薪资表格）
        """
        # 简化版结构化提取
        # 实际应用中应调用AI API进行智能提取
        
        lines = content.split('\n')
        structured = {
            "sections": [],
            "tables": [],
            "keywords": self._extract_keywords(content),
            "entities": []
        }
        
        current_section = None
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 简单的章节识别（实际应使用AI）
            if len(line) < 100 and (line.isupper() or line.startswith('第')):
                if current_section:
                    structured["sections"].append(current_section)
                current_section = {"title": line, "content": []}
            elif current_section:
                current_section["content"].append(line)
            else:
                if not structured["sections"]:
                    structured["sections"].append({"title": "正文", "content": [line]})
                else:
                    structured["sections"][-1]["content"].append(line)
        
        if current_section:
            structured["sections"].append(current_section)
        
        return structured
    
    def _extract_keywords(self, content: str, top_n: int = 20) -> List[str]:
        """提取关键词（简化版，实际应使用NLP模型）"""
        # 简单的关键词提取
        import re
        from collections import Counter
        
        # 移除标点符号，提取中文和英文单词
        words = re.findall(r'[\u4e00-\u9fff]+|[a-zA-Z]+', content.lower())
        
        # 过滤停用词（简化版）
        stop_words = {'的', '是', '在', '了', '和', '有', '为', '与', 'a', 'an', 'the', 'is', 'are', 'was', 'were'}
        words = [w for w in words if w not in stop_words and len(w) > 1]
        
        # 统计词频
        word_freq = Counter(words)
        return [word for word, _ in word_freq.most_common(top_n)]
    
    def save_parsed_document(self, filename: str, file_path: str, 
                            intent_space_id: Optional[int] = None) -> int:
        """
        保存解析后的文档到数据库
        
        Returns:
            文档ID
        """
        parse_result = self.parse_document(file_path, filename)
        
        conn = db.get_connection()
        cursor = conn.cursor()
        
        file_size = os.path.getsize(file_path)
        file_format = Path(filename).suffix.lower()
        
        cursor.execute("""
            INSERT INTO documents 
            (filename, file_path, file_format, file_size, status, intent_space_id, 
             processed_content, metadata)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            filename,
            file_path,
            file_format,
            file_size,
            "processed",
            intent_space_id,
            json.dumps(parse_result["structured_content"], ensure_ascii=False),
            json.dumps(parse_result["metadata"], ensure_ascii=False)
        ))
        
        doc_id = cursor.lastrowid
        conn.commit()
        conn.close()
        
        return doc_id


# 文档解析器实例
document_parser = DocumentParser()
